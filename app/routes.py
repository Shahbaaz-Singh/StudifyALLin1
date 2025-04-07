from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file, abort,make_response
from . import db, app
from .models import Admin, Employee, Customer, TravelHistory, WorkExperience, RefusalDetails,Education
from flask_login import login_user, logout_user, login_required, current_user
from .forms import CustomerForm, ChangePasswordForm
from datetime import datetime
from docx import Document
import os
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from functools import wraps

#--------------------------------------------------------------------------------------------------------------------------------#

main = Blueprint('main', __name__)

#--------------------------------------------------------------------------------------------------------------------------------#

def nocache(view):
    @wraps(view)
    def no_cache_view(*args, **kwargs):
        response = make_response(view(*args, **kwargs))
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response
    return no_cache_view

#--------------------------------------------------------------------------------------------------------------------------------#

@app.route('/')
@nocache
def index():
    return render_template('index.html')

#--------------------------------------------------------------------------------------#
# Admin Routes
#--------------------------------------------------------------------------------------#

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        admin = Admin.query.filter_by(username=username).first()

        if admin and admin.check_password(password):
            login_user(admin)
            return redirect(url_for('admin_dashboard'))
        else:
            flash('Invalid credentials, please try again', 'danger')
    
    return render_template('admin/admin_login.html')

@app.route('/admin/dashboard')
@login_required
@nocache
def admin_dashboard():
    return render_template('admin/admin_dashboard.html')

@app.route('/admin/change_password', methods=['GET', 'POST'])
@login_required
@nocache
def change_password_admin():
    form = ChangePasswordForm()
    admin = Admin.query.get(current_user.id)
    
    if not admin:
        flash("Employee not found.", "danger")
        return redirect(url_for('dashboard'))

    if form.validate_on_submit():
        if not admin.check_password(form.current_password.data):
            flash("Current password is incorrect.", "danger")
            return redirect(url_for('change_password_admin'))
        
        if form.current_password.data == form.new_password.data:
            flash("New password cannot be the same as the current password.", "danger")
            return redirect(url_for('change_password_admin'))
        
        if form.confirm_password.data != form.new_password.data:
            flash("Passwords should be same.", "danger")
            return redirect(url_for('change_password_admin'))
    
        admin.set_password(form.new_password.data)

        db.session.commit()

        flash("Your password has been successfully changed.", "success")
        return redirect(url_for('admin_dashboard'))

    return render_template('admin/change_password.html', form=form)

@app.route('/admin/manage_employees', methods=['GET'])
@login_required
@nocache
def manage_employees():
    employees = Employee.query.all()
    return render_template('admin/manage_employees.html', employees=employees)

@app.route('/admin/create_employee', methods=['GET', 'POST'])
@login_required
@nocache
def create_employee():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        phone_number = request.form.get('phone_number')
        email = request.form.get('email')
        address = request.form.get('address')

        new_employee = Employee(
            username=username,
            phone_number=phone_number,
            email=email,
            address=address,
            admin_id=current_user.id
        )
        new_employee.set_password(password)

        db.session.add(new_employee)
        db.session.commit()

        flash('Employee created successfully', 'success')
        return redirect(url_for('manage_employees'))

    return render_template('admin/create_employee.html')

@app.route('/admin/update_employee/<int:employee_id>', methods=['GET', 'POST'])
@login_required
@nocache
def update_employee(employee_id):
    employee = Employee.query.get_or_404(employee_id)

    if request.method == 'POST':
        username = request.form.get('username')
        phone_number = request.form.get('phone_number')
        email = request.form.get('email')
        address = request.form.get('address')

        employee.username = username
        employee.phone_number = phone_number
        employee.email = email
        employee.address = address

        db.session.commit()

        flash('Employee details updated successfully', 'success')
        return redirect(url_for('manage_employees')) 

    return render_template('admin/update_employee.html', employee=employee)

@app.route('/admin/manage_customers')
@login_required
@nocache
def manage_customers():
    return render_template('admin/manage_customers.html')

@app.route('/admin/create_customer', methods=['GET', 'POST'])
@login_required
@nocache
def create_customer_admin():
    form = CustomerForm()

    if form.validate_on_submit():
        try:
            customer = Customer(
                name=form.name.data,
                dob=form.dob.data,
                email=form.email.data,
                contact_no=form.contact_no.data,
                address=form.address.data,
                marital_status=form.marital_status.data,
                interested_in_ielts=form.interested_in_ielts.data,
                interested_in_pte=form.interested_in_pte.data,
                interested_in_toefl=form.interested_in_toefl.data,
                interested_in_others=form.interested_in_others.data,
                study_visa=form.study_visa.data,
                tourist_visa=form.tourist_visa.data,
                pr=form.pr.data,
                others_visa=form.others_visa.data,
                course_preference_1=form.course_preference_1.data,
                course_preference_2=form.course_preference_2.data,
                course_preference_3=form.course_preference_3.data,
                country_preference_1=form.country_preference_1.data,
                country_preference_2=form.country_preference_2.data,
                country_preference_3=form.country_preference_3.data,
                city_preference_1=form.city_preference_1.data,
                city_preference_2=form.city_preference_2.data,
                city_preference_3=form.city_preference_3.data,
                how_did_you_know_us=form.how_did_you_know_us.data,
                other_information=form.other_information.data,
                sign_date=datetime.utcnow(),
                admin_id=current_user.id,
                employee_id=current_user.id,
            )

            for education_form in form.education.entries:
                education = Education(
                    class_name=education_form.data['class_name'],
                    pass_out_year=education_form.data['pass_out_year'],
                    board_university_college=education_form.data['board_university_college'],
                    percentage_or_grade_point=education_form.data['percentage_or_grade_point'],
                    stream_subjects=education_form.data['stream_subjects'],
                )
                customer.education.append(education)

            for work_form in form.work_experience.entries:
                work_experience = WorkExperience(
                    organization_name=work_form.data['organization_name'],
                    designation=work_form.data['designation'],
                    joining_date=work_form.data['joining_date'],
                    leaving_date=work_form.data['leaving_date'],
                    full_time_part_time=work_form.data['full_time_part_time'],
                    job_duties=work_form.data['job_duties'],
                    was_it_your_business=work_form.data['was_it_your_business'],
                )
                customer.work_experience.append(work_experience)

            for refusal_form in form.refusal_details.entries:
                refusal_details = RefusalDetails(
                    country=refusal_form.data['country'],
                    refusal_date=refusal_form.data['refusal_date'],
                    reason=refusal_form.data['reason'],
                    additional_info=refusal_form.data['additional_info'],
                )
                customer.refusal_details.append(refusal_details)

            for travel_form in form.travel_history.entries:
                travel_history = TravelHistory(
                    country=travel_form.data['country'],
                    arrival_date=travel_form.data['arrival_date'],
                    departure_date=travel_form.data['departure_date'],
                    details=travel_form.data['details'],
                )
                customer.travel_history.append(travel_history)
                
                try:
                    db.session.add(customer)
                    db.session.commit()
                    flash('Applicant created successfully', 'success')
                except Exception as e:
                    db.session.rollback()

            return redirect(url_for('admin_dashboard'))

        except Exception as e:
            db.session.rollback()
            flash(f'Error creating customer: {e}', 'danger')

    return render_template('admin/create_customer.html', form=form)

@app.route('/admin/view_customers')
@login_required
@nocache
def view_customers_admin():

    customers = Customer.query.filter_by(admin_id=current_user.id).all()

    return render_template('admin/view_customers.html', customers=customers)


@app.route('/admin/view_customer/<int:customer_id>')
@login_required
@nocache
def view_details_admin(customer_id):
    customer = Customer.query.get_or_404(customer_id)
    
    if customer.admin_id != current_user.id:
        flash('You are not authorized to view this customer.', 'danger')
        return redirect(url_for('view_customers'))
    
    return render_template('admin/view_details.html', customer=customer)

@app.route('/admin/update_customer/<int:customer_id>', methods=['GET', 'POST'])
@login_required
@nocache
def update_customer_admin(customer_id):
    customer = Customer.query.get_or_404(customer_id)
    form = CustomerForm(obj=customer)

    if form.validate_on_submit():
        try:
            customer.name = form.name.data
            customer.dob = form.dob.data
            customer.email = form.email.data
            customer.contact_no = form.contact_no.data
            customer.address = form.address.data
            customer.marital_status = form.marital_status.data
            customer.interested_in_ielts = form.interested_in_ielts.data
            customer.interested_in_pte = form.interested_in_pte.data
            customer.interested_in_toefl = form.interested_in_toefl.data
            customer.interested_in_others = form.interested_in_others.data
            customer.study_visa = form.study_visa.data
            customer.tourist_visa = form.tourist_visa.data
            customer.pr = form.pr.data
            customer.others_visa = form.others_visa.data
            customer.course_preference_1 = form.course_preference_1.data
            customer.course_preference_2 = form.course_preference_2.data
            customer.course_preference_3 = form.course_preference_3.data
            customer.country_preference_1 = form.country_preference_1.data
            customer.country_preference_2 = form.country_preference_2.data
            customer.country_preference_3 = form.country_preference_3.data
            customer.city_preference_1 = form.city_preference_1.data
            customer.city_preference_2 = form.city_preference_2.data
            customer.city_preference_3 = form.city_preference_3.data
            customer.how_did_you_know_us = form.how_did_you_know_us.data
            customer.other_information = form.other_information.data

            for education_form in form.education.entries:
                education_id = education_form.data.get('id')
                if education_id:
                    education = Education.query.get(education_id)
                    if education:
                        education.class_name = education_form.data['class_name']
                        education.pass_out_year = education_form.data['pass_out_year']
                        education.board_university_college = education_form.data['board_university_college']
                        education.percentage_or_grade_point = education_form.data['percentage_or_grade_point']
                        education.stream_subjects = education_form.data['stream_subjects']
                else:
                    education = Education(
                        class_name=education_form.data['class_name'],
                        pass_out_year=education_form.data['pass_out_year'],
                        board_university_college=education_form.data['board_university_college'],
                        percentage_or_grade_point=education_form.data['percentage_or_grade_point'],
                        stream_subjects=education_form.data['stream_subjects'],
                    )
                    customer.education.append(education)

            for education in customer.education:
                if education.id not in [education_form.data.get('id') for education_form in form.education.entries]:
                    db.session.delete(education)

            for work_form in form.work_experience.entries:
                work_id = work_form.data.get('id')
                if work_id:
                    work_experience = WorkExperience.query.get(work_id)
                    if work_experience:
                        work_experience.organization_name = work_form.data['organization_name']
                        work_experience.designation = work_form.data['designation']
                        work_experience.joining_date = work_form.data['joining_date']
                        work_experience.leaving_date = work_form.data['leaving_date']
                        work_experience.full_time_part_time = work_form.data['full_time_part_time']
                        work_experience.job_duties = work_form.data['job_duties']
                        work_experience.was_it_your_business = work_form.data['was_it_your_business']
                else:
                    work_experience = WorkExperience(
                        organization_name=work_form.data['organization_name'],
                        designation=work_form.data['designation'],
                        joining_date=work_form.data['joining_date'],
                        leaving_date=work_form.data['leaving_date'],
                        full_time_part_time=work_form.data['full_time_part_time'],
                        job_duties=work_form.data['job_duties'],
                        was_it_your_business=work_form.data['was_it_your_business'],
                    )
                    customer.work_experience.append(work_experience)

            for work_experience in customer.work_experience:
                if work_experience.id not in [work_form.data.get('id') for work_form in form.work_experience.entries]:
                    db.session.delete(work_experience)

            for refusal_form in form.refusal_details.entries:
                refusal_id = refusal_form.data.get('id')
                if refusal_id:
                    refusal_details = RefusalDetails.query.get(refusal_id)
                    if refusal_details:
                        refusal_details.country = refusal_form.data['country']
                        refusal_details.refusal_date = refusal_form.data['refusal_date']
                        refusal_details.reason = refusal_form.data['reason']
                        refusal_details.additional_info = refusal_form.data['additional_info']
                else:
                    refusal_details = RefusalDetails(
                        country=refusal_form.data['country'],
                        refusal_date=refusal_form.data['refusal_date'],
                        reason=refusal_form.data['reason'],
                        additional_info=refusal_form.data['additional_info'],
                    )
                    customer.refusal_details.append(refusal_details)

            for refusal_details in customer.refusal_details:
                if refusal_details.id not in [refusal_form.data.get('id') for refusal_form in form.refusal_details.entries]:
                    db.session.delete(refusal_details)

            for travel_form in form.travel_history.entries:
                travel_id = travel_form.data.get('id')
                if travel_id:
                    travel_history = TravelHistory.query.get(travel_id)
                    if travel_history:
                        travel_history.country = travel_form.data['country']
                        travel_history.arrival_date = travel_form.data['arrival_date']
                        travel_history.departure_date = travel_form.data['departure_date']
                        travel_history.details = travel_form.data['details']
                else:
                    travel_history = TravelHistory(
                        country=travel_form.data['country'],
                        arrival_date=travel_form.data['arrival_date'],
                        departure_date=travel_form.data['departure_date'],
                        details=travel_form.data['details'],
                    )
                    customer.travel_history.append(travel_history)

            for travel_history in customer.travel_history:
                if travel_history.id not in [travel_form.data.get('id') for travel_form in form.travel_history.entries]:
                    db.session.delete(travel_history)

            db.session.commit()
            flash('Applicant updated successfully', 'success')
            return redirect(url_for('view_customers_admin'))

        except Exception as e:
            db.session.rollback()
            flash(f'Error updating customer: {e}', 'danger')

    return render_template('admin/create_customer.html', form=form)


@app.route('/admin/logout')
def admin_logout():
    logout_user()
    return redirect(url_for('admin_login'))

#----------------------------------------------------------------------------------------------------------------------#

#---------------------------------------------------------------------------------------------#
# Employee Routes
#---------------------------------------------------------------------------------------------#

@app.route('/employee/login', methods=['GET', 'POST'])
def employee_login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        employee = Employee.query.filter_by(username=username).first()

        if employee and employee.check_password(password):
            login_user(employee)
            return redirect(url_for('employee_dashboard'))
        else:
            flash('Invalid credentials, please try again', 'danger')
    
    return render_template('employee/employee_login.html')

@app.route('/employee/dashboard')
@login_required
@nocache
def employee_dashboard():
    return render_template('employee/employee_dashboard.html')

@app.route('/employee/change_password', methods=['GET', 'POST'])
@login_required
@nocache
def change_password():
    form = ChangePasswordForm()
    employee = Employee.query.get(current_user.id)
    
    if not employee:
        flash("Employee not found.", "danger")
        return redirect(url_for('dashboard'))

    if form.validate_on_submit():
        if not employee.check_password(form.current_password.data):
            flash("Current password is incorrect.", "danger")
            return redirect(url_for('change_password'))
        
        if form.current_password.data == form.new_password.data:
            flash("New password cannot be the same as the current password.", "danger")
            return redirect(url_for('change_password'))
        
        if form.confirm_password.data != form.new_password.data:
            flash("Passwords should be same.", "danger")
            return redirect(url_for('change_password'))
    
        employee.set_password(form.new_password.data)

        db.session.commit()

        flash("Your password has been successfully changed.", "success")
        return redirect(url_for('employee_dashboard'))

    return render_template('employee/change_password.html', form=form)

@app.route('/employee/create_customer', methods=['GET', 'POST'])
@login_required
@nocache
def create_customer():
    form = CustomerForm()

    if form.validate_on_submit():
        try:
            admin_id = current_user.admin_id
            customer = Customer(
                name=form.name.data,
                dob=form.dob.data,
                email=form.email.data,
                contact_no=form.contact_no.data,
                address=form.address.data,
                marital_status=form.marital_status.data,
                interested_in_ielts=form.interested_in_ielts.data,
                interested_in_pte=form.interested_in_pte.data,
                interested_in_toefl=form.interested_in_toefl.data,
                interested_in_others=form.interested_in_others.data,
                study_visa=form.study_visa.data,
                tourist_visa=form.tourist_visa.data,
                pr=form.pr.data,
                others_visa=form.others_visa.data,
                course_preference_1=form.course_preference_1.data,
                course_preference_2=form.course_preference_2.data,
                course_preference_3=form.course_preference_3.data,
                country_preference_1=form.country_preference_1.data,
                country_preference_2=form.country_preference_2.data,
                country_preference_3=form.country_preference_3.data,
                city_preference_1=form.city_preference_1.data,
                city_preference_2=form.city_preference_2.data,
                city_preference_3=form.city_preference_3.data,
                how_did_you_know_us=form.how_did_you_know_us.data,
                other_information=form.other_information.data,
                sign_date=datetime.utcnow(),
                admin_id=admin_id,
                employee_id=current_user.id,
            )

            for education_form in form.education.entries:
                education = Education(
                    class_name=education_form.data['class_name'],
                    pass_out_year=education_form.data['pass_out_year'],
                    board_university_college=education_form.data['board_university_college'],
                    percentage_or_grade_point=education_form.data['percentage_or_grade_point'],
                    stream_subjects=education_form.data['stream_subjects'],
                )
                customer.education.append(education)

            for work_form in form.work_experience.entries:
                work_experience = WorkExperience(
                    organization_name=work_form.data['organization_name'],
                    designation=work_form.data['designation'],
                    joining_date=work_form.data['joining_date'],
                    leaving_date=work_form.data['leaving_date'],
                    full_time_part_time=work_form.data['full_time_part_time'],
                    job_duties=work_form.data['job_duties'],
                    was_it_your_business=work_form.data['was_it_your_business'],
                )
                customer.work_experience.append(work_experience)

            for refusal_form in form.refusal_details.entries:
                refusal_details = RefusalDetails(
                    country=refusal_form.data['country'],
                    refusal_date=refusal_form.data['refusal_date'],
                    reason=refusal_form.data['reason'],
                    additional_info=refusal_form.data['additional_info'],
                )
                customer.refusal_details.append(refusal_details)

            for travel_form in form.travel_history.entries:
                travel_history = TravelHistory(
                    country=travel_form.data['country'],
                    arrival_date=travel_form.data['arrival_date'],
                    departure_date=travel_form.data['departure_date'],
                    details=travel_form.data['details'],
                )
                customer.travel_history.append(travel_history)
                
                try:
                    db.session.add(customer)
                    db.session.commit()
                    flash('Applicant created successfully', 'success')
                except Exception as e:
                    db.session.rollback()
                    
            return redirect(url_for('employee_dashboard'))

        except Exception as e:
            db.session.rollback()
            flash(f'Error creating customer: {e}', 'danger')

    return render_template('employee/create_customer.html', form=form)

@app.route('/employee/view_customers')
@login_required
@nocache
def view_customers():

    customers = Customer.query.filter_by(employee_id=current_user.id).all()

    return render_template('employee/view_customers.html', customers=customers)

@app.route('/employee/view_customer/<int:customer_id>')
@login_required
@nocache
def view_details(customer_id):
    customer = Customer.query.get_or_404(customer_id)
    
    if customer.employee_id != current_user.id:
        flash('You are not authorized to view this customer.', 'danger')
        return redirect(url_for('view_customers'))
    
    return render_template('employee/view_details.html', customer=customer)

@app.route('/employee/update_customer/<int:customer_id>', methods=['GET', 'POST'])
@login_required
@nocache
def update_customer(customer_id):
    customer = Customer.query.get_or_404(customer_id)
    form = CustomerForm(obj=customer)

    if form.validate_on_submit():
        try:
            customer.name = form.name.data
            customer.dob = form.dob.data
            customer.email = form.email.data
            customer.contact_no = form.contact_no.data
            customer.address = form.address.data
            customer.marital_status = form.marital_status.data
            customer.interested_in_ielts = form.interested_in_ielts.data
            customer.interested_in_pte = form.interested_in_pte.data
            customer.interested_in_toefl = form.interested_in_toefl.data
            customer.interested_in_others = form.interested_in_others.data
            customer.study_visa = form.study_visa.data
            customer.tourist_visa = form.tourist_visa.data
            customer.pr = form.pr.data
            customer.others_visa = form.others_visa.data
            customer.course_preference_1 = form.course_preference_1.data
            customer.course_preference_2 = form.course_preference_2.data
            customer.course_preference_3 = form.course_preference_3.data
            customer.country_preference_1 = form.country_preference_1.data
            customer.country_preference_2 = form.country_preference_2.data
            customer.country_preference_3 = form.country_preference_3.data
            customer.city_preference_1 = form.city_preference_1.data
            customer.city_preference_2 = form.city_preference_2.data
            customer.city_preference_3 = form.city_preference_3.data
            customer.how_did_you_know_us = form.how_did_you_know_us.data
            customer.other_information = form.other_information.data

            for education_form in form.education.entries:
                education_id = education_form.data.get('id')
                if education_id:
                    education = Education.query.get(education_id)
                    if education:
                        education.class_name = education_form.data['class_name']
                        education.pass_out_year = education_form.data['pass_out_year']
                        education.board_university_college = education_form.data['board_university_college']
                        education.percentage_or_grade_point = education_form.data['percentage_or_grade_point']
                        education.stream_subjects = education_form.data['stream_subjects']
                else:
                    education = Education(
                        class_name=education_form.data['class_name'],
                        pass_out_year=education_form.data['pass_out_year'],
                        board_university_college=education_form.data['board_university_college'],
                        percentage_or_grade_point=education_form.data['percentage_or_grade_point'],
                        stream_subjects=education_form.data['stream_subjects'],
                    )
                    customer.education.append(education)

            for education in customer.education:
                if education.id not in [education_form.data.get('id') for education_form in form.education.entries]:
                    db.session.delete(education)

            for work_form in form.work_experience.entries:
                work_id = work_form.data.get('id')
                if work_id:
                    work_experience = WorkExperience.query.get(work_id)
                    if work_experience:
                        work_experience.organization_name = work_form.data['organization_name']
                        work_experience.designation = work_form.data['designation']
                        work_experience.joining_date = work_form.data['joining_date']
                        work_experience.leaving_date = work_form.data['leaving_date']
                        work_experience.full_time_part_time = work_form.data['full_time_part_time']
                        work_experience.job_duties = work_form.data['job_duties']
                        work_experience.was_it_your_business = work_form.data['was_it_your_business']
                else:
                    work_experience = WorkExperience(
                        organization_name=work_form.data['organization_name'],
                        designation=work_form.data['designation'],
                        joining_date=work_form.data['joining_date'],
                        leaving_date=work_form.data['leaving_date'],
                        full_time_part_time=work_form.data['full_time_part_time'],
                        job_duties=work_form.data['job_duties'],
                        was_it_your_business=work_form.data['was_it_your_business'],
                    )
                    customer.work_experience.append(work_experience)

            for work_experience in customer.work_experience:
                if work_experience.id not in [work_form.data.get('id') for work_form in form.work_experience.entries]:
                    db.session.delete(work_experience)

            for refusal_form in form.refusal_details.entries:
                refusal_id = refusal_form.data.get('id')
                if refusal_id:
                    refusal_details = RefusalDetails.query.get(refusal_id)
                    if refusal_details:
                        refusal_details.country = refusal_form.data['country']
                        refusal_details.refusal_date = refusal_form.data['refusal_date']
                        refusal_details.reason = refusal_form.data['reason']
                        refusal_details.additional_info = refusal_form.data['additional_info']
                else:
                    refusal_details = RefusalDetails(
                        country=refusal_form.data['country'],
                        refusal_date=refusal_form.data['refusal_date'],
                        reason=refusal_form.data['reason'],
                        additional_info=refusal_form.data['additional_info'],
                    )
                    customer.refusal_details.append(refusal_details)

            for refusal_details in customer.refusal_details:
                if refusal_details.id not in [refusal_form.data.get('id') for refusal_form in form.refusal_details.entries]:
                    db.session.delete(refusal_details)

            for travel_form in form.travel_history.entries:
                travel_id = travel_form.data.get('id')
                if travel_id:
                    travel_history = TravelHistory.query.get(travel_id)
                    if travel_history:
                        travel_history.country = travel_form.data['country']
                        travel_history.arrival_date = travel_form.data['arrival_date']
                        travel_history.departure_date = travel_form.data['departure_date']
                        travel_history.details = travel_form.data['details']
                else:
                    travel_history = TravelHistory(
                        country=travel_form.data['country'],
                        arrival_date=travel_form.data['arrival_date'],
                        departure_date=travel_form.data['departure_date'],
                        details=travel_form.data['details'],
                    )
                    customer.travel_history.append(travel_history)

            for travel_history in customer.travel_history:
                if travel_history.id not in [travel_form.data.get('id') for travel_form in form.travel_history.entries]:
                    db.session.delete(travel_history)

            db.session.commit()
            flash('Applicant updated successfully', 'success')
            return redirect(url_for('view_customers'))

        except Exception as e:
            db.session.rollback()
            flash(f'Error updating customer: {e}', 'danger')

    return render_template('employee/create_customer.html', form=form)

@app.route('/employee/logout')
def employee_logout():
    logout_user()
    return redirect(url_for('employee_login'))

#------------------------------------------------------------------------------------------------------------------------------#

#--------------------------------------------------------------------------------------#
# Document Routes
#--------------------------------------------------------------------------------------#

def generate_customer_word(customer):
    doc_dir = os.path.join(os.getcwd(), 'static', 'word_docs')
    if not os.path.exists(doc_dir):
        os.makedirs(doc_dir)

    doc = Document()
    doc.add_heading(f"Customer Details: {customer.name}", 0)
    doc.add_paragraph(f"Customer ID: {customer.id}")
    doc.add_paragraph(f"Employee ID: {customer.employee_id}")
    doc.add_paragraph(f"Employee Name: {customer.employee.username}")
    doc.add_paragraph(f"Admin ID: {customer.admin_id}")
    doc.add_paragraph(f"Admin Name: {customer.admin.username}")

    doc.add_heading('Personal Information', level=1)
    doc.add_paragraph(f"Name: {customer.name}")
    doc.add_paragraph(f"Date of Birth: {customer.dob}")
    doc.add_paragraph(f"Email: {customer.email}")
    doc.add_paragraph(f"Contact Number: {customer.contact_no}")
    doc.add_paragraph(f"Address: {customer.address}")
    doc.add_paragraph(f"Marital Status: {customer.marital_status}")
    
    doc.add_heading('Visa & Immigration Information', level=1)
    doc.add_paragraph(f"Interested in IELTS: {'Yes' if customer.interested_in_ielts else 'No'}")
    doc.add_paragraph(f"Interested in PTE: {'Yes' if customer.interested_in_pte else 'No'}")
    doc.add_paragraph(f"Interested in TOEFL: {'Yes' if customer.interested_in_toefl else 'No'}")
    doc.add_paragraph(f"Other Interests: {customer.interested_in_others}")
    doc.add_paragraph(f"Study Visa: {'Yes' if customer.study_visa else 'No'}")
    doc.add_paragraph(f"Tourist Visa: {'Yes' if customer.tourist_visa else 'No'}")
    doc.add_paragraph(f"Permanent Residency: {'Yes' if customer.pr else 'No'}")
    doc.add_paragraph(f"Other Visa: {'Yes' if customer.others_visa else 'No'}")

    doc.add_heading('Course and Country Preferences', level=1)
    doc.add_paragraph(f"Course Preferences: {customer.course_preference_1}, {customer.course_preference_2}, {customer.course_preference_3}")
    doc.add_paragraph(f"Country Preferences: {customer.country_preference_1}, {customer.country_preference_2}, {customer.country_preference_3}")
    doc.add_paragraph(f"City Preferences: {customer.city_preference_1}, {customer.city_preference_2}, {customer.city_preference_3}")
    
    doc.add_heading('Travel History', level=1)
    for travel in customer.travel_history:
        doc.add_paragraph(f"Country: {travel.country}, Arrival Date: {travel.arrival_date}, Departure Date: {travel.departure_date}")
        doc.add_paragraph(f"Details: {travel.details}")

    doc.add_heading('Refusal Details', level=1)
    for refusal in customer.refusal_details:
        doc.add_paragraph(f"Country: {refusal.country}, Refusal Date: {refusal.refusal_date}, Reason: {refusal.reason}")
        doc.add_paragraph(f"Additional Info: {refusal.additional_info}")
    
    doc.add_heading('Work Experience', level=1)
    for work in customer.work_experience:
        doc.add_paragraph(f"Organization: {work.organization_name}, Designation: {work.designation}, Joining Date: {work.joining_date}, Leaving Date: {work.leaving_date}")
        doc.add_paragraph(f"Job Duties: {work.job_duties}")
        doc.add_paragraph(f"Was it your business: {work.was_it_your_business}")

    doc.add_heading('Education', level=1)
    for education in customer.education:
        doc.add_paragraph(f"Class: {education.class_name}, Year: {education.pass_out_year}, Board/University: {education.board_university_college}, Percentage: {education.percentage_or_grade_point}")
        doc.add_paragraph(f"Stream/Subjects: {education.stream_subjects}")

    doc.add_heading('How did you know us?', level=1)
    doc.add_paragraph(f"Source: {customer.how_did_you_know_us}")
    
    doc.add_heading('Other Information', level=1)
    doc.add_paragraph(f"Other Info: {customer.other_information}")
    
    doc.add_heading('Signature', level=1)
    doc.add_paragraph(f"Signature Date: {customer.sign_date.strftime('%Y-%m-%d') if customer.sign_date else 'N/A'}")

    doc_path = os.path.join(doc_dir, f"customer_{customer.id}_details.docx")
    doc.save(doc_path)

    return doc_path

@app.route('/customer/<int:customer_id>/download_word')
def download_word(customer_id):

    customer = Customer.query.get_or_404(customer_id)

    doc_path = generate_customer_word(customer)
    
    if not os.path.exists(doc_path):
        abort(404, description="Word document not found")
    
    return send_file(doc_path, as_attachment=True)

def generate_customer_pdf(customer):

    pdf_dir = os.path.join(os.getcwd(), 'static', 'pdf')
    if not os.path.exists(pdf_dir):
        os.makedirs(pdf_dir)

    pdf_path = os.path.join(pdf_dir, f"customer_{customer.id}_details.pdf")
    
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle('Heading1', parent=styles['Heading1'], fontSize=16, spaceAfter=12)
    section_style = ParagraphStyle('SectionHeader', parent=styles['Heading2'], fontSize=14, spaceAfter=8, textColor=colors.darkblue)
    detail_style = ParagraphStyle('DetailText', parent=styles['Normal'], fontSize=12, spaceAfter=4)

    doc = SimpleDocTemplate(pdf_path, pagesize=letter)

    content = []

    content.append(Paragraph(f"<b>Customer Details: {customer.name}</b>", heading_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph(f"<b>Customer ID: {customer.id}</b>", detail_style))
    content.append(Paragraph(f"<b>Employee ID: {customer.employee_id}</b>", detail_style))
    content.append(Paragraph(f"<b>Employee Name: {customer.employee.username}</b>", detail_style))
    content.append(Paragraph(f"<b>Admin ID: {customer.admin_id}</b>", detail_style))
    content.append(Paragraph(f"<b>Admin Name: {customer.admin.username}</b>", detail_style))

    content.append(Paragraph("<b>Personal Information</b>", section_style))
    content.append(Paragraph(f"Name: {customer.name}", detail_style))
    content.append(Paragraph(f"Date of Birth: {customer.dob.strftime('%Y-%m-%d')}", detail_style))
    content.append(Paragraph(f"Email: {customer.email}", detail_style))
    content.append(Paragraph(f"Contact Number: {customer.contact_no}", detail_style))
    content.append(Paragraph(f"Address: {customer.address}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Marital Status</b>", section_style))
    content.append(Paragraph(f"Status: {customer.marital_status}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Visa & Immigration Information</b>", section_style))
    content.append(Paragraph(f"Interested in IELTS: {'Yes' if customer.interested_in_ielts else 'No'}", detail_style))
    content.append(Paragraph(f"Interested in PTE: {'Yes' if customer.interested_in_pte else 'No'}", detail_style))
    content.append(Paragraph(f"Interested in TOEFL: {'Yes' if customer.interested_in_toefl else 'No'}", detail_style))
    content.append(Paragraph(f"Interested in Other: {'Yes' if customer.interested_in_others else 'No'}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Visa Preferences</b>", section_style))
    content.append(Paragraph(f"Study Visa: {'Yes' if customer.study_visa else 'No'}", detail_style))
    content.append(Paragraph(f"Tourist Visa: {'Yes' if customer.tourist_visa else 'No'}", detail_style))
    content.append(Paragraph(f"Permanent Residency: {'Yes' if customer.pr else 'No'}", detail_style))
    content.append(Paragraph(f"Other Visa: {'Yes' if customer.others_visa else 'No'}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Course Preferences</b>", section_style))
    content.append(Paragraph(f"Course Preference 1: {customer.course_preference_1}", detail_style))
    content.append(Paragraph(f"Course Preference 2: {customer.course_preference_2}", detail_style))
    content.append(Paragraph(f"Course Preference 3: {customer.course_preference_3}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Country & City Preferences</b>", section_style))
    content.append(Paragraph(f"Country 1: {customer.country_preference_1}", detail_style))
    content.append(Paragraph(f"Country 2: {customer.country_preference_2}", detail_style))
    content.append(Paragraph(f"Country 3: {customer.country_preference_3}", detail_style))
    content.append(Paragraph(f"City 1: {customer.city_preference_1}", detail_style))
    content.append(Paragraph(f"City 2: {customer.city_preference_2}", detail_style))
    content.append(Paragraph(f"City 3: {customer.city_preference_3}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Travel History</b>", section_style))
    for travel in customer.travel_history:
        arrival_date = travel.arrival_date.strftime('%Y-%m-%d') if travel.arrival_date else "N/A"
        departure_date = travel.departure_date.strftime('%Y-%m-%d') if travel.departure_date else "N/A"
        content.append(Paragraph(f"Country: {travel.country}, Arrival Date: {arrival_date}, Departure Date: {departure_date}", detail_style))
        content.append(Paragraph(f"Details: {travel.details if travel.details else 'N/A'}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Refusal Details</b>", section_style))
    for refusal in customer.refusal_details:
        refusal_date = refusal.refusal_date.strftime('%Y-%m-%d') if refusal.refusal_date else "N/A"
        content.append(Paragraph(f"Country: {refusal.country}, Refusal Date: {refusal_date}, Reason: {refusal.reason}", detail_style))
        content.append(Paragraph(f"Additional Info: {refusal.additional_info if refusal.additional_info else 'N/A'}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Work Experience</b>", section_style))
    for work in customer.work_experience:
        joining_date = work.joining_date.strftime('%Y-%m-%d') if work.joining_date else "N/A"
        leaving_date = work.leaving_date.strftime('%Y-%m-%d') if work.leaving_date else "N/A"
        content.append(Paragraph(f"Organization: {work.organization_name}, Designation: {work.designation}, Joining Date: {joining_date}, Leaving Date: {leaving_date}", detail_style))
        content.append(Paragraph(f"Job Duties: {work.job_duties if work.job_duties else 'N/A'}", detail_style))
        content.append(Paragraph(f"Was it your business: {work.was_it_your_business if work.was_it_your_business else 'N/A'}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Education Information</b>", section_style))
    for education in customer.education:
        content.append(Paragraph(f"Class: {education.class_name}, Pass Out Year: {education.pass_out_year}, College: {education.board_university_college}, Grade: {education.percentage_or_grade_point}", detail_style))
        content.append(Paragraph(f"Stream/Subjects: {education.stream_subjects if education.stream_subjects else 'N/A'}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Additional Information</b>", section_style))
    content.append(Paragraph(f"Other Information: {customer.other_information}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>How did you know us?</b>", section_style))
    content.append(Paragraph(f"How did you know us: {customer.how_did_you_know_us}", detail_style))
    content.append(Spacer(1, 12))

    content.append(Paragraph("<b>Signature</b>", section_style))
    content.append(Paragraph(f"Signature Date: {customer.sign_date.strftime('%Y-%m-%d') if customer.sign_date else 'N/A'}", detail_style))

    doc.build(content)

    return pdf_path

@app.route('/customer/<int:customer_id>/download_pdf')
def download_pdf(customer_id):
    customer = Customer.query.get_or_404(customer_id)

    pdf_path = generate_customer_pdf(customer)
    
    if not os.path.exists(pdf_path):
        abort(404, description="PDF document not found")
    
    return send_file(pdf_path, as_attachment=True)

#--------------------------------------------------------------------------------------------------------------#